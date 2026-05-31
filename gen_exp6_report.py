"""
软件工程实验六报告生成脚本
实验六 面向对象软件案例的UML建模
"""
import os
import sys
sys.path.insert(0, r'F:\SoftwareEngineering\.claude\skills\se-exp-report\references')
from report_utils import *
from docx import Document

TEMPLATE = r'F:\SoftwareEngineering\计算机科学与工程学院实验报告模板.docx'
OUT = r'F:\SoftwareEngineering\softwareproject\experiment6\实验六_面向对象软件案例的UML建模.docx'
SS = r'F:\SoftwareEngineering\softwareproject\experiment6\screenshots'


def fill_header(table):
    # Row 0: 学号 + 姓名
    clear_cell(table.rows[0].cells[1])
    add_text_paragraph(table.rows[0].cells[1], '202331607106',
                      align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    clear_cell(table.rows[0].cells[5])
    add_text_paragraph(table.rows[0].cells[5], '韩志鑫',
                      align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    # Row 1: 专业 + 班级
    clear_cell(table.rows[1].cells[1])
    add_text_paragraph(table.rows[1].cells[1], '计算机科学与技术',
                      align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    clear_cell(table.rows[1].cells[5])
    add_text_paragraph(table.rows[1].cells[5], '23级卓越班',
                      align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    # Row 2: 课程名称 + 课程类型
    clear_cell(table.rows[2].cells[1])
    add_text_paragraph(table.rows[2].cells[1], '软件工程',
                      align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    clear_cell(table.rows[2].cells[5])
    add_text_paragraph(table.rows[2].cells[5], '必修课',
                      align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    # Row 3: 实验名称
    clear_cell(table.rows[3].cells[1])
    add_text_paragraph(table.rows[3].cells[1], '实验六 面向对象软件案例的UML建模',
                      align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)


def fill_purpose(table):
    cell = table.rows[4].cells[0]
    clear_cell_keep_label(cell)
    add_text_paragraph(cell, '（1）练习并掌握JSP+JavaBean+Servlet技术实现的MVC设计模式；')
    add_text_paragraph(cell, '（2）练习UML建模技术；')
    add_text_paragraph(cell, '（3）实验任务采用多人协作完成。')


def fill_content(table):
    cell = table.rows[5].cells[0]
    clear_cell_keep_label(cell)

    # ═══ 任务1 ═══
    add_heading_paragraph(cell, '一、任务1：创建学生信息管理的Java Web项目')

    add_subheading_paragraph(cell, '1.1 创建MySQL数据库及表（完成人：韩志鑫）')
    add_text_paragraph(cell, '在MySQL中创建数据库students，并在students数据库中创建学生信息表student。student表包含7个字段：id（主键自增）、studentNo（学号）、name（姓名）、gender（性别）、age（年龄）、major（专业）、className（班级），并插入5条测试数据。')
    add_table_caption(cell, '表1 student表结构')
    add_three_line_table(cell,
        headers=['字段名', '数据类型', '约束', '说明'],
        rows=[
            ['id', 'INT', 'PRIMARY KEY AUTO_INCREMENT', '学生ID'],
            ['studentNo', 'VARCHAR(20)', 'NOT NULL', '学号'],
            ['name', 'VARCHAR(50)', 'NOT NULL', '姓名'],
            ['gender', 'VARCHAR(4)', "DEFAULT '男'", '性别'],
            ['age', 'INT', '', '年龄'],
            ['major', 'VARCHAR(100)', '', '专业'],
            ['className', 'VARCHAR(50)', '', '班级'],
        ],
        col_widths_cm=[2.5, 2.5, 4, 3.5]
    )
    add_image(cell, os.path.join(SS, '截图一.png'), width_cm=13, height_cm=8)
    add_image_caption(cell, '图1 MySQL数据库students及student表数据')

    add_subheading_paragraph(cell, '1.2 创建Web项目与包结构（完成人：韩志鑫）')
    add_text_paragraph(cell, '在Eclipse中创建Dynamic Web Project，项目名称为studentmanager。在项目src包中创建包model、control、entity、dbutil，在WebRoot下创建jsp文件夹，在WEB-INF下配置web.xml。项目采用标准的Java Web目录结构，遵循MVC设计模式。')
    add_image(cell, os.path.join(SS, '截图2.png'), width_cm=13, height_cm=8)
    add_image_caption(cell, '图2 Eclipse项目目录结构')

    add_subheading_paragraph(cell, '1.3 编写实体类与数据库工具类（完成人：韩志鑫）')
    add_text_paragraph(cell, '在entity包中创建Student.java实体类，属性与数据库表student字段一一对应，包含无参和全参构造方法以及getter/setter方法。在dbutil包中创建Dbconn.java数据库连接工具类，封装获取Connection连接的静态方法getConnection()和关闭资源的close()方法，使用com.mysql.cj.jdbc.Driver驱动连接MySQL数据库。')

    add_subheading_paragraph(cell, '1.4 编写Model层与Controller层（完成人：韩志鑫）')
    add_text_paragraph(cell, '在model包中创建StudentModel.java学生信息处理类，实现getAllStudents()、getStudentById()、insertStudent()、updateStudent()、deleteStudent()、insertStudentBatch()六个方法，分别对应查询所有、按ID查询、新增、修改、删除、批量导入功能，所有数据库操作均使用PreparedStatement防止SQL注入。')
    add_text_paragraph(cell, '在control包中创建6个Servlet控制器：ListStudentServlet（列表）、InsertStudentServlet（新增）、UpStudentServlet（修改前查询）、DoStudentServlet（修改后更新）、ShowStudentServlet（删除前确认）、DeleteStudentServlet（删除后执行），并在web.xml中配置servlet和servlet-mapping。')

    add_subheading_paragraph(cell, '1.5 编写JSP视图层（完成人：韩志鑫）')
    add_text_paragraph(cell, '在jsp文件夹下创建4个JSP视图页面：studentlist.jsp（学生列表）、studentinsert.jsp（新增学生）、studentupdate.jsp（修改学生）、studentshow.jsp（删除确认）。所有页面通过Servlet控制器转发访问，避免直接访问JSP导致的相对路径不一致问题。同时创建index.jsp作为系统入口，自动重定向到ListStudentServlet。')

    add_subheading_paragraph(cell, '1.6 项目部署与功能测试（完成人：韩志鑫、肖辉京）')
    add_text_paragraph(cell, '将项目部署到Apache Tomcat 10服务器。由于Tomcat 10基于Jakarta EE 9，需要将所有Servlet的javax.servlet导入改为jakarta.servlet，并将web.xml的schema版本设为5.0。部署完成后，通过浏览器访问http://localhost:8080/studentmanager/，测试学生信息的增删改查操作。')

    add_image(cell, os.path.join(SS, '截图3-首页.png'), width_cm=13, height_cm=8)
    add_image_caption(cell, '图3 学生信息列表页面')

    add_image(cell, os.path.join(SS, '截图4-新增学生信息页.png'), width_cm=13, height_cm=8)
    add_image_caption(cell, '图4 新增学生信息页面')

    add_image(cell, os.path.join(SS, '截图4-新增学生信息成功.png'), width_cm=13, height_cm=8)
    add_image_caption(cell, '图5 新增学生信息成功')

    # ═══ 任务2 ═══
    add_heading_paragraph(cell, '二、任务2：学生信息批量导入功能增量开发')
    add_subheading_paragraph(cell, '2.1 CSV批量导入功能开发（完成人：韩志鑫）')
    add_text_paragraph(cell, '为案例软件新增CSV文件批量导入学生信息的功能。创建ImportStudentServlet.java控制器和studentimport.jsp视图页面。用户上传CSV格式文件（逗号分隔，UTF-8编码），系统解析每一行数据并批量调用insertStudentBatch()方法写入数据库。CSV文件可由Excel直接编辑后另存为CSV格式，兼容性好且无需额外第三方依赖。')
    add_text_paragraph(cell, '同时在web.xml中为ImportStudentServlet配置multipart-config，设置文件大小阈值、最大上传大小等参数，支持大文件上传。')

    add_image(cell, os.path.join(SS, '截图5-批量导入学生信息成功.png'), width_cm=13, height_cm=8)
    add_image_caption(cell, '图6 CSV批量导入学生信息')

    add_image(cell, os.path.join(SS, '截图5-批量导入学生信息成功后的主页.png'), width_cm=13, height_cm=8)
    add_image_caption(cell, '图7 批量导入后的学生列表')

    # ═══ 任务3 ═══
    add_heading_paragraph(cell, '三、任务3：GitHub仓库创建与源码上传')
    add_subheading_paragraph(cell, '3.1 项目仓库（完成人：韩志鑫）')
    add_text_paragraph(cell, '在本人GitHub账号下创建项目仓库studentmanager，上传实验完成的全部软件源码。仓库包含完整的项目源码（entity/dbutil/model/control四个包、JSP视图页面、CSS样式表）、数据库建表脚本（student.sql）、项目配置文件（web.xml、.classpath、.project等）、依赖JAR包、UML图源文件、README说明文档。')
    add_text_paragraph(cell, 'GitHub仓库地址：https://github.com/Billhzx/studentmanager')

    # ═══ 任务4 ═══
    add_heading_paragraph(cell, '四、任务4：学生信息管理系统类图')
    add_subheading_paragraph(cell, '4.1 类图绘制（完成人：韩志鑫）')
    add_text_paragraph(cell, '利用在线作图工具绘制学生信息管理软件案例的类图。类图包含以下核心类：Student（学生实体类）、Dbconn（数据库连接工具类）、StudentModel（学生信息模型类）、以及7个Servlet控制器类（ListStudentServlet、InsertStudentServlet、UpStudentServlet、DoStudentServlet、ShowStudentServlet、DeleteStudentServlet、ImportStudentServlet），均继承自HttpServlet。类之间的关系包括：Servlet类依赖StudentModel进行业务逻辑处理，StudentModel依赖Dbconn获取数据库连接、依赖Student传递数据，Servlet类与Student之间存在使用关系。')

    add_image(cell, os.path.join(SS, '系统类图.png'), width_cm=13, height_cm=9)
    add_image_caption(cell, '图8 学生信息管理系统类图')

    # ═══ 任务5 ═══
    add_heading_paragraph(cell, '五、任务5：学生信息管理系统用例图')
    add_subheading_paragraph(cell, '5.1 用例图绘制（完成人：韩志鑫）')
    add_text_paragraph(cell, '利用在线作图工具绘制学生信息管理软件的用例图。系统参与者为系统管理员，用例包括：查看学生列表、新增学生信息、修改学生信息、删除学生信息、CSV批量导入。其中修改和删除用例扩展自查询学生详情用例。')

    add_image(cell, os.path.join(SS, '用例图.png'), width_cm=13, height_cm=9)
    add_image_caption(cell, '图9 学生信息管理系统用例图')

    # ═══ 任务6 ═══
    add_heading_paragraph(cell, '六、任务6：实验六WBS编制')
    add_subheading_paragraph(cell, '6.1 WBS工作分解（完成人：韩志鑫）')
    add_text_paragraph(cell, '采用自顶向下的方法，将实验六的整体工作分解为7个一级任务（创建学生信息管理Java Web项目、Excel/CSV批量导入增量开发、GitHub仓库创建与源码上传、绘制类图、绘制用例图、编制WBS、撰写实验报告），再进一步细化为23个二级工作包。WBS活动耗时如下表所示。')

    add_table_caption(cell, '表2 实验六WBS活动耗时表')
    add_three_line_table(cell,
        headers=['编号', '工作包', '活动内容', '预估(分)', '实际(分)', '完成人'],
        rows=[
            ['1.1', '数据库', '创建MySQL数据库及student表', '10', '10', '韩志鑫'],
            ['1.2', '项目搭建', '创建Web项目及包结构', '10', '10', '韩志鑫'],
            ['1.3', '环境配置', '添加MySQL驱动包', '5', '5', '韩志鑫'],
            ['1.4', '实体层', '编写Student.java', '10', '10', '韩志鑫'],
            ['1.5', '工具层', '编写Dbconn.java', '10', '10', '韩志鑫'],
            ['1.6', '模型层', '编写StudentModel.java', '20', '20', '韩志鑫'],
            ['1.7', '控制层', '编写6个Servlet控制器', '30', '30', '韩志鑫'],
            ['1.8', '视图层', '编写4个JSP页面+CSS样式', '30', '30', '韩志鑫'],
            ['1.9', '配置', '配置web.xml及Eclipse项目文件', '15', '15', '韩志鑫'],
            ['1.10', '调试', 'Tomcat 10适配与部署调试', '30', '30', '韩志鑫'],
            ['1.11', '测试', 'CRUD功能测试', '15', '15', '韩志鑫/肖辉京'],
            ['2.1', '增量开发', '编写ImportStudentServlet及JSP页面', '20', '20', '韩志鑫'],
            ['2.2', '增量开发', 'CSV导入功能测试', '5', '5', '韩志鑫'],
            ['3.1', 'GitHub', '创建仓库并上传项目源码', '10', '10', '韩志鑫'],
            ['4.1', 'UML建模', '分析并绘制类图', '15', '15', '韩志鑫'],
            ['5.1', 'UML建模', '分析并绘制用例图', '10', '10', '韩志鑫'],
            ['6.1', 'WBS', '编制WBS及耗时估算', '15', '15', '韩志鑫'],
            ['7.1', '文档', '截图整理与过程陈述', '20', '20', '韩志鑫'],
            ['7.2', '文档', '撰写实验总结', '15', '15', '韩志鑫'],
            ['7.3', '文档', '生成报告文档', '10', '10', '韩志鑫'],
        ],
        col_widths_cm=[1.2, 2, 5.5, 1.5, 1.5, 2.8]
    )
    add_text_paragraph(cell, '合计预估耗时：315分钟；合计实际耗时：315分钟。小组成员分工：韩志鑫承担约90%的任务量，肖辉京承担约10%的任务量。')


def fill_summary(table):
    cell = table.rows[6].cells[0]
    clear_cell_keep_label(cell)

    add_subheading_paragraph(cell, '1. 小组成员名单')
    add_text_paragraph(cell, '韩志鑫（学号：202331607106）、肖辉京')

    add_subheading_paragraph(cell, '2. 我承担的工作')
    add_text_paragraph(cell, '在本次实验中，我（韩志鑫）承担了约90%的任务量，主要负责以下工作：')
    add_multi_run_paragraph(cell, [
        ('（1）项目搭建与环境配置：', HEI, True),
        ('创建Java Web项目studentmanager，配置包结构（entity/dbutil/model/control），添加MySQL驱动包，配置Eclipse项目文件和Tomcat 10部署环境。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（2）MVC三层编码：', HEI, True),
        ('编写实体类Student.java、数据库工具类Dbconn.java、模型类StudentModel.java（含增删改查和批量导入方法），以及6个Servlet控制器和4个JSP视图页面，完整实现学生信息的CRUD功能。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（3）Tomcat 10适配：', HEI, True),
        ('排查并解决Tomcat 10环境下javax.servlet到jakarta.servlet的迁移问题，解决@WebServlet注解导致的类加载失败问题，统一所有页面通过Servlet转发访问。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（4）增量开发：', HEI, True),
        ('独立完成CSV批量导入功能的开发，包括ImportStudentServlet控制器和studentimport.jsp视图页面的编写。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（5）UML建模：', HEI, True),
        ('完成学生信息管理系统的类图和用例图设计与绘制。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（6）WBS编制与文档撰写：', HEI, True),
        ('完成WBS分解、耗时估算以及实验报告的编写。', SONG, False)
    ])

    add_subheading_paragraph(cell, '3. 实验体会与感受')
    add_text_paragraph(cell, '通过本次实验，我对JSP+JavaBean+Servlet的MVC设计模式有了深入的理解和体会：')
    add_multi_run_paragraph(cell, [
        ('（1）MVC分层架构的优势：', HEI, True),
        ('本次实验严格按照MVC模式组织代码：Student实体类对应Model数据模型层，StudentModel处理业务逻辑，Servlet作为Controller控制器协调请求转发，JSP作为View视图层负责页面展示，Dbconn提供数据访问支持。这种分层架构使得各模块职责清晰、耦合度低，便于维护和扩展。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（2）Tomcat 10的Jakarta EE迁移：', HEI, True),
        ('Tomcat 10基于Jakarta EE 9，命名空间从javax.servlet全面迁移到jakarta.servlet，web.xml的schema也需要升级到5.0版本。此外，@WebServlet注解在Tomcat 10的注解扫描过程中可能触发类加载失败，最终改用web.xml配置servlet解决了该问题。这些踩坑经历让我深刻认识到Web服务器版本升级带来的兼容性影响。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（3）Servlet路径统一的重要性：', HEI, True),
        ('在开发过程中遇到了因直接访问JSP和通过Servlet转发混合使用导致的相对路径混乱问题。最终采用"所有页面统一通过Servlet转发访问"的策略，在需要展示页面的Servlet中添加doGet方法转发到对应JSP，彻底解决了路径问题。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（4）CSV优于POI的实践选择：', HEI, True),
        ('批量导入功能最初计划使用Apache POI读取Excel文件，但POI依赖链过长（需要xmlbeans、commons-collections4、commons-compress、commons-io、log4j-api等多个jar包），且出现了多次ClassNotFoundException。最终改用CSV格式实现，零额外依赖，Excel可直接编辑后另存为CSV，简单可靠。', SONG, False)
    ])
    add_multi_run_paragraph(cell, [
        ('（5）团队协作的体会：', HEI, True),
        ('本次实验以我为主进行开发，肖辉京同学参与功能测试和部分截图工作。虽然分工不均衡，但整个协作过程沟通顺畅、效率较高。', SONG, False)
    ])


def clear_commentary(table):
    cell = table.rows[7].cells[0]
    clear_cell_keep_label(cell)


def main():
    doc = Document(TEMPLATE)
    table = doc.tables[0]
    fill_header(table)
    fill_purpose(table)
    fill_content(table)
    fill_summary(table)
    clear_commentary(table)
    doc.save(OUT)
    print(f'已保存到 {OUT}')


if __name__ == '__main__':
    main()
